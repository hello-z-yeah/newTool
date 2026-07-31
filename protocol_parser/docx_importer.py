"""从 Word 协议文档自动生成 JSON 配置。

支持的文档结构：
- 表格型：识别包含 "命令字"/"cmd"/"attrid" 等关键词的表格
- 段落型：提取关键信息（帧头、版本、校验等）
- 混合型：自动适配

使用：
    from protocol_parser.docx_importer import import_from_docx
    cfg = import_from_docx("协议.docx")
    # 保存为 JSON
    import json
    with open("product/my_product.json", "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class ImporterError(Exception):
    """文档导入错误。"""


# ---------- 工具函数 ----------

def _normalize(s: str) -> str:
    """标准化字符串：去除空白、换行、特殊字符。"""
    if not s:
        return ""
    return re.sub(r"\s+", " ", s).strip()


def _find_int(s: str) -> int | None:
    """从字符串中提取整数（支持 0x 前缀）。"""
    if not s:
        return None
    s = s.strip()
    # 优先匹配 0x 开头
    m = re.search(r"0[xX]([0-9a-fA-F]+)", s)
    if m:
        return int(m.group(1), 16)
    # 匹配纯数字
    m = re.search(r"\b(\d+)\b", s)
    if m:
        return int(m.group(1))
    return None


def _find_hex_int(s: str) -> int | None:
    """从字符串中提取 hex 整数（支持 0A / 0E / 00 这种无前缀形式）。"""
    if not s:
        return None
    s = s.strip()
    # 0x 前缀
    m = re.search(r"0[xX]([0-9a-fA-F]+)", s)
    if m:
        return int(m.group(1), 16)
    # 纯 hex 字符串（1-8 位，含字母才算 hex，避免与十进制冲突）
    m = re.match(r"^[0-9a-fA-F]{1,8}$", s)
    if m:
        return int(s, 16)
    # 字符串开头的 hex 字段
    m = re.match(r"^([0-9a-fA-F]{1,8})\b", s)
    if m:
        candidate = m.group(1)
        # 必须含字母才认定是 hex（纯数字走十进制）
        if any(c in "abcdefABCDEF" for c in candidate):
            return int(candidate, 16)
        return int(candidate)
    return None


def _find_hex_bytes(s: str) -> str | None:
    """从字符串中提取 hex 字节序列，返回标准格式 '0xXXXX'。"""
    if not s:
        return None
    # 匹配 0xA5A5 / 0xA5 0xA5 / A5 A5 A5
    m = re.search(r"0[xX]([0-9a-fA-F]{2,4})", s)
    if m:
        return f"0x{m.group(1).upper()}"
    # 匹配 A5A5 这种连续 hex
    m = re.search(r"\b([0-9a-fA-F]{4})\b", s)
    if m:
        return f"0x{m.group(1).upper()}"
    return None


def _table_to_rows(table) -> list[list[str]]:
    """把 docx 表格转成二维字符串列表。"""
    rows = []
    for row in table.rows:
        cells = [_normalize(c.text) for c in row.cells]
        rows.append(cells)
    return rows


def _find_header_row(rows: list[list[str]], keywords: list[str]) -> int:
    """在表格中找表头行（包含所有关键词的行）。"""
    for i, row in enumerate(rows):
        row_text = " ".join(row).lower()
        if all(k.lower() in row_text for k in keywords):
            return i
    return -1


def _column_map(header: list[str], field_names: list[list[str]]) -> dict[str, int]:
    """根据表头建立字段名 → 列索引的映射。

    field_names 是 [[可能名1, 可能名2], ...] 形式。
    """
    mapping: dict[str, int] = {}
    header_lower = [h.lower() for h in header]
    for field, names in field_names.items():
        for i, h in enumerate(header_lower):
            if any(n.lower() in h for n in names):
                mapping[field] = i
                break
    return mapping


# ---------- 文档结构识别 ----------

@dataclass
class ParsedDocument:
    """从 Word 解析出的中间结构。"""
    product_name: str = ""
    description: str = ""
    frame_config: dict = field(default_factory=dict)
    commands: list[dict] = field(default_factory=list)
    attributes: dict = field(default_factory=dict)
    enums: dict = field(default_factory=dict)
    raw_tables: list[list[list[str]]] = field(default_factory=list)
    raw_paragraphs: list[str] = field(default_factory=list)


def _read_docx(path: str | Path) -> ParsedDocument:
    """读取 Word 文档，提取所有段落和表格。"""
    if not HAS_DOCX:
        raise ImporterError(
            "python-docx 未安装。请在命令行执行：\n"
            "  pip install python-docx -i https://pypi.tuna.tsinghua.edu.cn/simple"
        )
    p = Path(path)
    if not p.exists():
        raise ImporterError(f"文件不存在: {p}")

    doc = Document(p)
    parsed = ParsedDocument()

    # 段落
    for para in doc.paragraphs:
        text = _normalize(para.text)
        if text:
            parsed.raw_paragraphs.append(text)

    # 表格
    for table in doc.tables:
        rows = _table_to_rows(table)
        parsed.raw_tables.append(rows)

    # 从段落提取产品名/描述
    for text in parsed.raw_paragraphs[:20]:
        if any(k in text.lower() for k in ["协议", "protocol"]):
            if not parsed.product_name:
                parsed.product_name = text[:60]
                continue
        if not parsed.description and len(text) > 10:
            parsed.description = text[:200]
            break

    return parsed


# ---------- 帧结构识别 ----------

def _parse_frame_config(parsed: ParsedDocument) -> dict:
    """从文档识别帧结构配置。"""
    frame = {
        "header": "0xA5A5",
        "header_size": 2,
        "ver": "0x03",
        "ver_offset": 2,
        "ver_size": 1,
        "cmd_offset": 3,
        "length_offset": 4,
        "length_size": 2,
        "length_byte_order": "big",
        "checksum": {
            "algorithm": "sum",
            "length": 1,
            "covers": "from_start_to_checksum_exclusive",
        },
    }

    # 从段落中找帧头、版本、校验等关键词
    all_text = "\n".join(parsed.raw_paragraphs)
    for table in parsed.raw_tables:
        for row in table:
            all_text += "\n" + " ".join(row)

    # 帧头：寻找 "帧头 0xA5A5" / "起始符 0xA5" 等
    m = re.search(r"(?:帧头|起始符|header|frame header)[^\n]*?(0[xX][0-9a-fA-F]{2,4})", all_text, re.IGNORECASE)
    if m:
        val = m.group(1)
        frame["header"] = val
        hex_str = val[2:]
        frame["header_size"] = len(hex_str) // 2

    # 版本
    m = re.search(r"(?:版本|version|ver)[^\n]*?(0[xX][0-9a-fA-F]+)", all_text, re.IGNORECASE)
    if m:
        frame["ver"] = m.group(1)

    # 校验算法
    if re.search(r"异或|XOR", all_text, re.IGNORECASE):
        frame["checksum"]["algorithm"] = "xor"
    elif re.search(r" CRC", all_text, re.IGNORECASE):
        frame["checksum"]["algorithm"] = "crc8"
    elif re.search(r"求和|sum|累加", all_text, re.IGNORECASE):
        frame["checksum"]["algorithm"] = "sum"

    return frame


# ---------- 命令表识别 ----------

CMD_FIELD_NAMES = {
    "cmd_code": ["cmd", "命令字", "命令码", "命令", "数值", "cmd_code", "value"],
    "name": ["name", "名称", "命令名", "功能", "描述", "description"],
    "description": ["说明", "描述", "功能描述", "备注", "description"],
    "direction": ["方向", "direction", "类型"],
    "format": ["格式", "format", "数据格式"],
}


def _parse_commands(parsed: ParsedDocument) -> list[dict]:
    """识别命令字表。

    策略：找到含 "命令字"/"cmd"/"数值" 等表头的表格，按列解析。
    """
    commands: list[dict] = []
    seen_codes: set[int] = set()

    for table in parsed.raw_tables:
        if not table:
            continue
        # 找表头行
        header_idx = _find_header_row(table, ["命令字"])
        if header_idx < 0:
            header_idx = _find_header_row(table, ["cmd"])
        if header_idx < 0:
            # V3.0 风格：Name + 方向 + 数值 + 说明
            for i, row in enumerate(table):
                row_text = " ".join(row).lower()
                if ("数值" in row_text or "value" in row_text) and ("name" in row_text or "名称" in row_text or "说明" in row_text):
                    header_idx = i
                    break
        if header_idx < 0:
            # 也尝试 "命令" + "名称" 组合
            for i, row in enumerate(table):
                row_text = " ".join(row).lower()
                if ("命令" in row_text or "cmd" in row_text) and ("名称" in row_text or "name" in row_text or "功能" in row_text):
                    header_idx = i
                    break
        if header_idx < 0:
            continue

        header = table[header_idx]
        col_map = _column_map(header, CMD_FIELD_NAMES)
        if "cmd_code" not in col_map:
            continue

        cmd_col = col_map["cmd_code"]
        name_col = col_map.get("name")
        desc_col = col_map.get("description", name_col)
        dir_col = col_map.get("direction")
        fmt_col = col_map.get("format")

        for row in table[header_idx + 1:]:
            if cmd_col >= len(row):
                continue
            cmd_text = row[cmd_col]
            # 命令字通常以 hex 形式书写（0x20 / 20 / 0xA5 等），优先按 hex 解析
            cmd_code = _find_hex_int(cmd_text)
            if cmd_code is None:
                cmd_code = _find_int(cmd_text)
            if cmd_code is None or cmd_code in seen_codes:
                continue

            name = row[name_col] if name_col is not None and name_col < len(row) else ""
            desc = row[desc_col] if desc_col is not None and desc_col < len(row) else ""
            direction = row[dir_col] if dir_col is not None and dir_col < len(row) else ""
            fmt = row[fmt_col] if fmt_col is not None and fmt_col < len(row) else ""

            # 推断 format
            if not fmt:
                fmt = _infer_format_from_text(name + " " + desc)

            cmd = {
                "cmd_code": f"0x{cmd_code:02X}",
                "name": name or f"cmd_{cmd_code:02X}",
                "description": desc,
            }
            if direction:
                cmd["direction"] = direction

            # 双向命令（同 cmd_code 有 request/response）
            # 简化处理：根据方向字段分拆；无方向字段则默认 request/response 都用 attr_list
            if direction.lower() in ("请求", "request", "下行", "命令下发"):
                cmd["request"] = {"format": fmt or "raw", "name": name}
                cmd["response"] = {"format": "raw", "name": name + " 响应"}
            elif direction.lower() in ("响应", "response", "上行", "应答"):
                # 已存在同 cmd_code 的 request 则补充；否则当作双向
                existing = next((c for c in commands if c.get("cmd_code") == f"0x{cmd_code:02X}"), None)
                if existing and "request" in existing:
                    existing["response"] = {"format": fmt or "raw", "name": name}
                    continue
                else:
                    cmd["request"] = {"format": "raw", "name": name + " 请求"}
                    cmd["response"] = {"format": fmt or "raw", "name": name}
            else:
                # 无方向：默认双向
                cmd["request"] = {"format": "raw", "name": "请求"}
                cmd["response"] = {"format": fmt or "attr_list", "name": "响应"}

            commands.append(cmd)
            seen_codes.add(cmd_code)

    return commands


def _infer_format_from_text(text: str) -> str:
    """根据命令描述推断数据格式。"""
    text_lower = text.lower()
    if "心跳" in text or "heartbeat" in text_lower:
        return "module_status"
    if "设备信息" in text or "version" in text_lower or "dev_info" in text_lower:
        return "dev_version"
    if "配网" in text or "net_config" in text_lower:
        return "net_config"
    if "时间" in text and "get" in text_lower:
        return "get_time_resp"
    if "快照" in text or "snapshot" in text_lower:
        return "attr_list"
    if "上报" in text or "report" in text_lower:
        return "attr_list"
    if "事件" in text or "event" in text_lower:
        return "event"
    if "行为" in text or "action" in text_lower:
        return "msg_id_then_action"
    if "ota" in text_lower or "升级" in text:
        return "raw"
    if "错误" in text or "errcode" in text_lower:
        return "errcode"
    return "attr_list"


# ---------- 属性表识别 ----------

ATTR_FIELD_NAMES = {
    "attrid": ["attrid", "attr id", "属性id", "属性id", "属性", "aid"],
    # 英文/系统属性名：Name列（on / night-light-switch / mode / fan-level ...）
    "name": ["name", "属性名"],
    # 中文属性名称列（照明 / 人感夜灯 / 模式 / 吹风档位 ...）单独用 cn_name 保存，不覆盖 name
    "cn_name": ["属性名称", "中文名称", "中文名", "名称"],
    "typeid": ["typeid", "type id", "类型id", "类型", "type"],
    "access": ["access", "权限", "访问", "读写"],
    "unit": ["unit", "单位"],
    "range": ["range", "范围", "取值"],
    "enum": ["enum", "枚举", "取值说明", "说明"],
}


TYPEID_KEYWORDS = {
    "bool": 0, "布尔": 0,
    "int8": 1,
    "uint8": 2, "u8": 2,
    "int16": 3, "i16": 3,
    "uint16": 4, "u16": 4,
    "int32": 5, "i32": 5,
    "uint32": 6, "u32": 6,
    "int64": 7,
    "uint64": 8,
    "float": 9, "float32": 9,
    "string": 11, "字符串": 11,
    "f1_u16": 15, "f1u16": 15,
    "f2_u16": 16, "f2u16": 16,
}


def _parse_typeid(text: str) -> int | None:
    """从文本解析 typeid。"""
    if not text:
        return None
    text_lower = text.strip().lower()
    # 直接是数字
    n = _find_int(text)
    if n is not None and 0 <= n <= 24:
        return n
    # 关键词匹配
    for kw, tid in TYPEID_KEYWORDS.items():
        if kw in text_lower:
            return tid
    return None


def _parse_attributes(parsed: ParsedDocument) -> dict:
    """识别属性表。

    支持两种格式：
    1. 表格型：包含 attrid/属性id 列的表格
    2. 段落型：如 "0x01   工作状态" + "类型: typeid=2, 访问: ?" + "枚举: ..." 的段落组
    """
    attributes: dict = {}

    # ---- 表格型属性表 ----
    for table in parsed.raw_tables:
        if not table:
            continue
        header_idx = -1
        for i, row in enumerate(table):
            row_text = " ".join(row).lower()
            if ("attrid" in row_text or "属性id" in row_text or "属性 id" in row_text) and ("name" in row_text or "名称" in row_text):
                header_idx = i
                break
        if header_idx < 0:
            continue

        header = table[header_idx]
        col_map = _column_map(header, ATTR_FIELD_NAMES)
        if "attrid" not in col_map:
            continue

        aid_col = col_map["attrid"]
        name_col = col_map.get("name")
        cn_name_col = col_map.get("cn_name")
        type_col = col_map.get("typeid")
        access_col = col_map.get("access")
        unit_col = col_map.get("unit")
        range_col = col_map.get("range")
        enum_col = col_map.get("enum")

        for row in table[header_idx + 1:]:
            if aid_col >= len(row):
                continue
            aid_text = row[aid_col]
            aid = _find_hex_int(aid_text)
            if aid is None:
                aid = _find_int(aid_text)
            if aid is None:
                continue

            key = f"0x{aid:02X}"
            if key in attributes:
                continue

            raw_name = row[name_col].strip() if (name_col is not None and name_col < len(row) and row[name_col]) else ""
            raw_cn_name = row[cn_name_col].strip() if (cn_name_col is not None and cn_name_col < len(row) and row[cn_name_col]) else ""

            # 容错：有些表格只提供一列「名称/属性名称」—— 要根据内容猜测到底是中文还是英文 Name
            if raw_name and (not raw_cn_name):
                # 没专门cn_name列，或name列与cn_name列索引相同→自动分拆：中文内容塞进 cn_name
                if any("\u4e00" <= ch <= "\u9fff" for ch in raw_name):
                    raw_cn_name, raw_name = raw_name, raw_cn_name
            if (not raw_cn_name) and name_col is not None and cn_name_col is not None and name_col == cn_name_col:
                # 两列映射到同一索引（如「名称」既匹配 name 又匹配 cn_name 关键字）→当只有一列时，有中文时当cn_name用
                pass
            attr: dict[str, Any] = {"name": raw_name}
            if raw_cn_name:
                attr["cn_name"] = raw_cn_name

            if type_col is not None and type_col < len(row):
                tid = _parse_typeid(row[type_col])
                if tid is not None:
                    attr["typeid"] = tid

            if access_col is not None and access_col < len(row):
                attr["access"] = row[access_col]

            if unit_col is not None and unit_col < len(row) and row[unit_col]:
                attr["unit"] = row[unit_col]

            if range_col is not None and range_col < len(row) and row[range_col]:
                attr["range"] = row[range_col]

            if enum_col is not None and enum_col < len(row) and row[enum_col]:
                enum_text = row[enum_col]
                enum_map = _parse_enum_text(enum_text)
                if enum_map:
                    attr["enum"] = enum_map

            attributes[key] = attr

    # ---- 段落型属性表 ----
    paragraph_attrs = _parse_attributes_from_paragraphs(parsed.raw_paragraphs)
    for key, attr in paragraph_attrs.items():
        if key not in attributes:
            attributes[key] = attr

    return attributes


def _parse_attributes_from_paragraphs(paragraphs: list[str]) -> dict:
    """从段落文本中识别属性表。

    匹配模式：
        0x01   工作状态
              类型: typeid=2, 访问: ?
              枚举:
                1: 待机中 2: 制冰中 3: 清洗中 4: 冰满
        0x02   故障
              类型: typeid=2, 访问: ?
              范围 无故障 缺水 ...
    """
    import re as _re
    attributes: dict = {}
    i = 0
    n = len(paragraphs)

    # 属性ID行正则
    attr_id_pattern = _re.compile(r'^\s*(0[xX][0-9a-fA-F]{1,2})\s+(.+)$')
    # 类型行
    type_pattern = _re.compile(r'^\s*(?:类型|type)[：:]\s*(.+?)\s*$', _re.IGNORECASE)
    # 枚举开始行
    enum_start_pattern = _re.compile(r'^\s*(?:枚举|enum)\s*[：:]\s*$', _re.IGNORECASE)
    # 范围行
    range_pattern = _re.compile(r'^\s*(?:范围|range)\s+(.+?)\s*$', _re.IGNORECASE)

    while i < n:
        line = paragraphs[i]
        m = attr_id_pattern.match(line)
        if not m:
            i += 1
            continue

        aid_hex = m.group(1)
        name = m.group(2).strip()
        try:
            aid = int(aid_hex, 16)
        except ValueError:
            i += 1
            continue

        i += 1
        attr: dict[str, Any] = {"name": name}

        enum_lines: list[str] = []
        in_enum = False

        while i < n:
            next_line = paragraphs[i]
            # 下一个属性ID行，退出
            if attr_id_pattern.match(next_line):
                break

            # 类型行
            tm = type_pattern.match(next_line)
            if tm:
                type_text = tm.group(1)
                tid = _parse_typeid(type_text)
                if tid is not None:
                    attr["typeid"] = tid
                acc_match = _re.search(r'(?:访问|access)\s*[=：:]\s*(\S+)', type_text, _re.IGNORECASE)
                if acc_match:
                    attr["access"] = acc_match.group(1)
                i += 1
                continue

            # 枚举开始行
            if enum_start_pattern.match(next_line):
                in_enum = True
                i += 1
                continue

            # 范围行
            rm = range_pattern.match(next_line)
            if rm:
                attr["range"] = rm.group(1).strip()
                i += 1
                continue

            # 枚举项
            if in_enum:
                enum_item_match = _re.match(r'^\s*(\d+)\s*[：:=]\s*(.+)', next_line)
                if enum_item_match:
                    enum_lines.append(next_line.strip())
                    i += 1
                    continue
                # 枚举结束
                in_enum = False
                if enum_lines:
                    enum_map = _parse_enum_text("; ".join(enum_lines))
                    if enum_map:
                        attr["enum"] = enum_map
                    enum_lines = []
                i += 1
                continue

            if next_line.strip() == "":
                i += 1
                continue

            break

        # 处理剩余枚举
        if enum_lines:
            enum_map = _parse_enum_text("; ".join(enum_lines))
            if enum_map:
                attr["enum"] = enum_map

        key = f"0x{aid:02X}"
        if key not in attributes and attr.get("name"):
            attributes[key] = attr

    return attributes


def _parse_enum_text(text: str) -> dict[str, str]:
    """解析枚举文本，兼容多种写法：
       - "0:关 1:开"
       - "0=关闭, 1=打开"
       - "1: 待机中 ↵ 2: 制冰中 ↵ 3: 清洗中" （多行/中文冒号/顿号）
       - "0-低档，1-高档" （短横线+全角逗号）
    同时会过滤「最小值:30 / 最大值:42 / 步长:1」这种纯范围描述，不放到枚举里。
    """
    if not text:
        return {}
    # 先分行，再用分隔符切，避免一行里范围说明和枚举混在一起
    result: dict[str, str] = {}
    # 过滤掉非枚举关键字前缀：最小值/最大值/步长/默认值/默认
    _FILTER_PREFIXES = ("最小值", "最大值", "最小", "最大", "步长", "默认值", "默认", "min", "max", "step", "default", "单位", "精度")
    # 先按所有常见分隔符拆
    raw_lines = re.split(r"[\r\n]+", text)
    pieces: list[str] = []
    for line in raw_lines:
        line = line.strip()
        if not line:
            continue
        sub_parts = re.split(r"[;,，；、]+", line)
        for sp in sub_parts:
            sp = sp.strip()
            if sp:
                pieces.append(sp)
    for part in pieces:
        # 如果整段就是范围描述，跳过整段（但不能误伤含有「最小值:X」前导 + 枚举项跟在同一行后面的情况，因此只是过滤每个 piece 的前缀 match，不是整行过滤）
        if any(re.match(rf"^\s*{re.escape(p)}[\s:：=].*$", part, re.IGNORECASE) for p in _FILTER_PREFIXES):
            continue
        # 匹配多种「数字 + 分隔符(冒号/等号/中文冒号/短横线) + 描述」
        matches = re.findall(r'(\d+)\s*[:=：\-]\s*(.*?)(?=\s+\d+\s*[:=：\-]|$)', part)
        for k, v in matches:
            v = v.strip().strip('"\'「」『』（）()[]【】')
            if not v:
                continue
            # 再过滤一次值前缀（以防 "最小值:30" 被正则后半截当 30:xxx 误抓）
            if any(v.startswith(p) for p in _FILTER_PREFIXES):
                continue
            if k not in result:
                result[k] = v
    return result


# ---------- 主入口 ----------

def import_from_docx(path: str | Path, product_name: str | None = None) -> dict:
    """从 Word 文档生成协议 JSON 配置。

    Args:
        path: Word 文档路径
        product_name: 自定义产品名（不指定则从文档提取）

    Returns:
        协议配置字典，可直接保存为 JSON。
        若导入过程中发现可疑问题（如命令列表为空、属性列表为空、无任何表头等），
        cfg["_import_warnings"] 会附带字符串列表，调用方（如 GUI）可以据此弹 warning。
    """
    warnings: list[str] = []
    parsed = _read_docx(path)

    total_tables = len(parsed.raw_tables)
    total_paragraphs = len(parsed.raw_paragraphs)
    # 文档内容为空（没读到任何段落/表格）→ 直接告警，避免用户以为导入成功实际什么都没解析
    if total_paragraphs == 0 and total_tables == 0:
        warnings.append("Word 文档未读取到任何段落或表格，请检查 docx 文件是否有效、是否加密、是否为 .doc 老格式。")

    # 进一步解析
    parsed.frame_config = _parse_frame_config(parsed)
    parsed.commands = _parse_commands(parsed)
    parsed.attributes = _parse_attributes(parsed)

    if not parsed.commands:
        warnings.append(
            "命令列表为空（未解析到任何命令）：请确认表格中包含「命令字/Name/说明/请求响应」列，"
            "且表头包含「cmd_code/命令/指令」等关键字。"
        )
    if not parsed.attributes:
        warnings.append(
            "属性列表为空（未解析到任何属性）：请确认属性表存在且表头包含「attrID/Name/属性名称/Type」列。"
        )
    if not parsed.frame_config:
        warnings.append(
            "帧配置为空：未在文档中发现帧头/长度/校验配置，将使用程序内置 V3.0 默认帧结构。"
        )
    if total_tables == 0 and (parsed.commands or parsed.attributes):
        warnings.append(
            "未检测到任何表格，但解析结果非空（可能全部来自正文字符串正则匹配）。"
            "建议使用标准表格填写协议，可显著提升导入准确性。"
        )

    # 产品名
    name = product_name or parsed.product_name or Path(path).stem
    # 清理产品名（只保留字母数字下划线）
    name = re.sub(r"[^\w\u4e00-\u9fff]", "_", name)
    name = re.sub(r"_+", "_", name).strip("_").lower() or "product"

    cfg = {
        "product": name,
        "description": parsed.description or f"从 {Path(path).name} 导入的协议",
        "version": "1.0",
        "frame": parsed.frame_config,
        "enums": parsed.enums,
        "commands": parsed.commands,
        "attributes": parsed.attributes,
        "_imported_from": str(Path(path).name),
    }
    if warnings:
        cfg["_import_warnings"] = warnings

    return cfg


def save_protocol_json(cfg: dict, output_path: str | Path) -> Path:
    """保存协议配置为 JSON 文件。"""
    p = Path(output_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    with p.open("w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)
    return p


def import_and_save(docx_path: str | Path, protocols_dir: str | Path, product_name: str | None = None) -> tuple[dict, Path]:
    """导入 Word 文档并保存为 JSON。

    Returns: (配置字典, 保存路径)
    """
    cfg = import_from_docx(docx_path, product_name)
    product = cfg["product"]
    out_path = Path(protocols_dir) / f"{product}.json"
    save_protocol_json(cfg, out_path)
    return cfg, out_path


def check_docx_available() -> bool:
    """检查 python-docx 是否可用。"""
    return HAS_DOCX
